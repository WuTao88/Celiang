import sys
import os
from docx import Document
from docx.shared import Pt
from docx.shared import Inches
from docx.oxml.ns import qn
import openpyxl as xl
from tkinter import *
from PIL import Image ,ImageTk
from tkinter import messagebox
from tkinter import ttk
import random
import math
from win32com import client as wc
import win32print
import tempfile
import win32api
import pythoncom
import CeLiang
import time
from natsort import natsorted
import threading

#################################################################

def source_path(relative_path):
    if os.path.exists(relative_path):
        base_path = os.path.abspath(".")
    elif getattr(sys, 'frozen', False):
        base_path = sys._MEIPASS
    return os.path.join(base_path, relative_path)


background=source_path('background')
RES=source_path('res')
muban=source_path('muban') 

################################################################
class OPTION:

    def getBM(zh=0.0,filename=None):
        global res

        name=res+'\\BM' if filename ==None else filename
        data=open(name, "r+",encoding='UTF-8')
        gc={}
        for sz in data:        
            gc=eval(sz)
        near=[abs(round(i-zh,3)) for i in gc.keys()]
        key1=round(zh+min(near),3)
        key2=round(zh-min(near),3)

        if (key1) in gc:
            return (key1,gc[key1][0],gc[key1][1])
        else:
            return (key2,gc[key2][0],gc[key2][1])


    def getCtr(zh=0.0,filename=None):
        global res
        name=res+'\\ContralPoint' if filename ==None else filename
        data=open(name, "r+",encoding='UTF-8')
        KZD={}
        for v in data:        
            KZD=eval(v)
        keys=list(KZD.keys())
        near=[abs(round(i-zh,3)) for i in KZD.keys()]
        key1=round(zh+min(near),3)
        key2=round(zh-min(near),3)
        a=None
        dt=[]
        if (key1) in KZD:
            a=key1
            dt.append(KZD[key1])
        else:
            a=key2
            dt.append(KZD[key2])
        if keys.index(a)+1 <len(keys):
            dt.append(KZD[keys[keys.index(a)+1]])
        else:
            dt.append(KZD[keys[keys.index(a)-1]])
        return dt

########################################################################

class RMC:
    def LuJian(zh):
        if zh>6100 and zh<9645.5:
            lj=0.4
        elif zh>9645.5:
            lj=0.5
        return lj
    def Widen_Road():
        Rw=6.5
        if zh>=9645.5:
            Rw=7.5
        return Rw

###############################


class MyThread(threading.Thread):
    def __init__(self, func, *args):
        super().__init__()

        
        self.func = func
        self.args = args
        
        self.setDaemon(True)
        self.start()    # 在这里开始
        
    def run(self):
        self.func(*self.args)



class SP:

    def AngleChange(angle):
        d = int(angle)
        m = int((angle - int(angle)) * 60)
        s = round((angle - int(angle)) * 3600 - m * 60, 2)
        dms = "{0}°{1}′{2}″".format(d, m, s)
        return dms


    def mileageToStr(km):
        k = int(km // 1000)
        b = int((km - k * 1000) // 100)
        s = int((km - k * 1000 - b * 100) // 10)
        g = int((km - k * 1000 - b * 100 - s * 10) // 1)
        f = km - km // 1
        ZH = "K{0}+{1}{2}{3:0>3.03f}".format(k, b, s, g + f)
        return ZH



    def LMZD(bm=[],H=0.0,L=30.0,data=[],i=0,Hsx=None):
        
        name=bm[1]
        Hsz=bm[2]
        hd=abs(H/L*30)
        if i==0 and Hsx==None:
            if hd <4 and hd>3.5:
                sz=round(random.uniform(hd+0.7,hd+1.0) if H<0 else random.uniform(0.7,1.0),3)
                dest=round(random.uniform(hd+0.7,hd+1.0) if H>0 else random.uniform(0.7,1.0),3)
            elif hd>=4:
                hd=2.0
                sz=round(random.uniform(hd+0.7,hd+1.0) if H<0 else random.uniform(0.7,1.0),3)
                dest=round(random.uniform(hd+0.7,hd+1.0) if H>0 else random.uniform(0.7,1.0),3)         
            else:
                sz=round(random.uniform(hd+1.0,hd+1.6) if H<0 else random.uniform(1.0,1.6),3)
                dest=round(random.uniform(hd+1.0,hd+1.6) if H>0 else random.uniform(1.0,1.6),3)
                hd=2.0
            if  abs(H)<4.1 and abs(L)<30:
                Hsx=round(Hsz+sz,3)
                data.append([name,'',sz,'','',Hsx,'',Hsz,'',''])
                return (Hsx,dest,1)
            else:
                data.append([name,'',sz,'','','','',Hsz,'',''])
            Hsx=round(Hsz+sz,3)
        else:
            sz=0
            dest=0
            hd=2.0 if hd>4.0 else hd
                    
        i=1 if i==0 else i
        while round(H+sz-dest,3)!=0:
            hs=round(random.uniform(0.7+hd,1.0+hd) if H<0 else random.uniform(0.7,1.0),3)
            qs=round(random.uniform(0.7,1.0) if H<0 else random.uniform(0.7+hd,1.0+hd),3)
            if abs(round(H+sz-dest,3))<=hd and i>=int(L//30):
                qs=round(hs+H+sz-dest,3)
                if qs<0.4:
                    qst=round(random.uniform(0.7,0.7+hd),3)
                    d=round(qst-qs,3)
                    hs=round(hs+d,3)
                    qs=qst
                elif qs>4.8:
                    qst=round(random.uniform(4.5-hd,4.5),3)
                    d=round(qst-qs,3)
                    hs=round(hs+d,3)
                    qs=qst                    
                Hsx=round(Hsx+hs-qs,3)
                data.append(['ZD%d'%i,'',hs,'',qs,Hsx,'','',''])
                break
            else:
                data.append(['ZD%d'%i,'',hs,'',qs,'','','',''])
                Hsx=round(Hsx+hs-qs,3)
                H=round(H+hs-qs,3)
                i=i+1                           
        return (Hsx,dest,i+1)

    def ZD(bm=[],H=0.0,L=30.0,data=[],i=0,Hsx=None):
        
        name=bm[1]
        Hsz=bm[2]
        hd=abs(H/L*30) 
        if i==0 and Hsx==None:
            if hd <4 and hd>3.5:
                sz=round(random.uniform(hd+0.7,hd+1.0) if H<0 else random.uniform(0.7,1.0),3)
                dest=round(random.uniform(hd+0.7,hd+1.0) if H>0 else random.uniform(0.7,1.0),3)
            elif hd>=4:
                hd=2.0
                sz=round(random.uniform(hd+0.7,hd+1.0) if H<0 else random.uniform(0.7,1.0),3)
                dest=round(random.uniform(hd+0.7,hd+1.0) if H>0 else random.uniform(0.7,1.0),3)         
            else:
                sz=round(random.uniform(hd+1.0,hd+1.6) if H<0 else random.uniform(1.0,1.6),3)
                dest=round(random.uniform(hd+1.0,hd+1.6) if H>0 else random.uniform(1.0,1.6),3)
                hd=2.0
            if  abs(H)<4.1 and abs(L)<30:
                Hsx=round(Hsz+sz,3)
                data.append([name,sz,'','',Hsx,'',Hsz,'',''])
                return (Hsx,dest,1)
            else:
                data.append([name,sz,'','','','',Hsz,'',''])
            Hsx=round(Hsz+sz,3)
        else:
            sz=0
            dest=0
            hd=2.0 if hd>4.0 else hd
            
        i=1 if i==0 else i
        while round(H+sz-dest,3)!=0:
            hs=round(random.uniform(0.7+hd,1.0+hd) if H<0 else random.uniform(0.7,1.0),3)
            qs=round(random.uniform(0.7,1.0) if H<0 else random.uniform(0.7+hd,1.0+hd),3)
            if abs(round(H+sz-dest,3))<=hd and i>=int(L//30):
                qs=round(hs+H+sz-dest,3)                
                if qs<0.4:
                    qst=round(random.uniform(0.7,0.7+hd),3)
                    d=round(qst-qs,3)
                    hs=round(hs+d,3)
                    qs=qst
                elif qs>4.8:
                    qst=round(random.uniform(4.5-hd,4.5),3)
                    d=round(qst-qs,3)
                    hs=round(hs+d,3)
                    qs=qst
                Hsx=round(Hsx+hs-qs,3)
                data.append(['ZD%d'%i,hs,'',qs,Hsx,'','',''])
                break
            else:
                data.append(['ZD%d'%i,hs,'',qs,'','','',''])
                Hsx=round(Hsx+hs-qs,3)
                H=round(H+hs-qs,3)
                i=i+1                           
        return (Hsx,dest,i+1)

    
    def HEADER(KZD=[]):
        
        if KZD!=None:
            X0=KZD[0][1]
            Y0=KZD[0][2]
            H0=KZD[0][3]
            X2=KZD[1][1]
            Y2=KZD[1][2]
            H2=KZD[1][3]
            s=round(((X2-X0)**2+(Y2-Y0)**2)**0.5,3)
            a=math.degrees(math.atan2(X2-X0,Y2-Y0))
            α=SP.AngleChange(a) if a>0 else SP.AngleChange(a+360)
            KZD.append(s)
            KZD.append(α)
            return KZD
        else:
            return None




    def DQSIZE(rhf=3.25,H0=5.0,mode='衡重式挡墙'):
        global res        
        h0=round(H0,2)
        file=open(res+'\\'+'DQ_size', "r+",encoding='UTF-8')
        dq={}
        for h in file:
            dq=eval(h)
        def getDQ(dq,mode,h0):
            return dq[mode][h0] if h0 in dq[mode] else None
        d=getDQ(dq,mode,h0)
        if d==None:
            return None
        #"H0":['h3','b4','b21','n','hn', 'm','m1']
        h3=d[0]
        b4=d[1]
        b21=d[2]
        n=d[3]
        hn=d[4]
        m=d[5]
        m1=d[6]
        H=round(h0-h3-hn,3)
        if mode !='路堑墙':
            ex1=round(hn*1.5,3)
            ex2=round(h3*m,3) if mode=='仰斜式挡墙' else 0
            jk1=round(rhf+0.5+H*m+b21+ex2-ex1,3)
            jk2=round(jk1+ex1-b4,3)
            jd1=round(rhf+0.5+H*m+b21+ex2,3)
            jd2=round(jd1-b4,3)
            jcd1=round(rhf+0.5+H*m+b21,3)
            jcd2=round(jd2-(hn+h3)*m1,3)
            qsd1=round(rhf+0.5,3)
            qsd2=rhf
        else:
            
            ex1=round(hn*1.5,3)
            jk1=round(rhf+0.4-0.5*m1-b21+ex1,3)
            jk2=round(rhf+0.4-0.5*m1-b21+b4,3)
            jd1=round(rhf+0.4-0.5*m1-b21,3)
            jd2=round(rhf+0.4-0.5*m1-b21+b4,3)
            jcd1=round(rhf+0.4-0.5*m1-b21,3)
            jcd2=round(rhf+0.4-0.5*m1-b21+b4,3)
            qsd1=round(rhf+0.4-0.5*m1+H*m1,3)
            qsd2=round(rhf+0.4-0.5*m1-b21+b4,3)
        return{'基坑开挖后':[jk1,jk2],'基坑底':[jd1,jd2],'基础顶':[jcd1,jcd2],'墙身顶':[qsd1,qsd2]}
            

    #衡重式挡墙
    #data=[reHigh,side,'仰斜式挡墙','基坑底']
    def DQ_PJ(zh,rw,LJ,data=[]):
        global res

        def getdq(dq):
            
            return ((data[1] in dq) and (zh>dq[0] and zh<dq[1]))
        file=open(res+'\\DQ_shoufang', "r+",encoding='UTF-8')

        dqs=[]
        for d in file:
            dqs=eval(d)
        
        chicun=list(filter(getdq,dqs))
       
        if chicun==[]:
            return None
        
        H=chicun[0][5]
        print(H)
        rh=abs(CeLiang.CeLiang(res,zh,rw).side(data[1])[1])
        print(rh)
        PJ= SP.DQSIZE(rh,H,data[2])[data[3]]

        pianju= [[zh,i] for i in PJ] if data[1]=='右侧' else [[zh,-i] for i in PJ]
        return pianju
        


    #圆管涵
    #data=[reHigh,d,'圆管涵','工序']
    def HD(zh,rw,LJ,data=[]):
        bd={0.75:0.11,1.0:0.12,1.5:0.14}
        δ=bd[data0]
        k1=round(1.5*data[0]+2*δ+0.3+0.5,3)

        return{'基坑开挖后':[round(zh-k1,3),round(zh+k1,3)]}



    #土方路基
    # zh 桩号，rw 道路宽度，LJ 路肩宽度，reHigh 相对设计路面高度。
    #data=[reHigh]
    def TF(zh,rw,LJ,data=[]):
        global res
        wide=CeLiang.CeLiang(res,zh,rw).widen
        ex=-data[0]*1.5
        return [[zh,round(wide[0]-ex,3)],[zh,0],[zh,round(wide[1]+ex,3)]]

    # zh 桩号，rw 道路宽度，LJ 路肩宽度，SIDE 路肩位置。
    #data=[reHigh,SIDE]
    def LJ(zh,rw,LJ,data=[]):
        global res
        Half=CeLiang.CeLiang(res,zh,rw).side(data[0])[1]
        ex=LJ if  data[0]=='右侧' else -LJ
        return [[zh,Half],[zh,round(Half-ex,3)]]

    # zh 桩号，rw 道路宽度，LJ 路肩宽度，SIDE 边沟位置，TP 边沟型号，reHigh 设计路面到边沟顶的高度。
    #data=[reHigh,side,TP]
    def BG(zh,rw,LJ,data=[]):
        global res
        Half=CeLiang.CeLiang(res,zh,rw).side(data[0])[1]

        if 'Ⅰ型' in data[2]:
            ex=0.65 if data[0]=='右侧' else -0.65
            return [[zh,Half],[zh,0],[zh,Half+ex]]
        elif 'Ⅱ型' in data[2]:
            ex=0.4 if data[0]=='右侧' else -0.4
            return [[zh,Half],[zh,Half+ex]]
        elif 'Ⅲ型' in data[2]:
            ex1=round(data[1]*1.5,3)
            ex2=round(ex1+0.9,3)
            pass
        elif 'Ⅳ型' in data[2]:
            ex1=-0.25 if data[0]=='右侧' else 0.25
            ex2= 0.75 if data[0]=='右侧' else -0.75
            return [[zh,round(Half-ex1,3)],[zh,round(Half+ex2,3)]]
    # zh 桩号，rw 道路宽度，LJ 路肩宽度
    #data=[reHigh]
    def JPSS(zh,rw,LJ,data=[]):
        global res
        wide=CeLiang.CeLiang(res,zh,rw).widen
        return [[zh,round(wide[0]+LJ,3)],[zh,0],[zh,round(wide[1]-LJ,3)]]



    # zh 桩号，rw 道路宽度，LJ 路肩宽度
    #data=[reHigh]
    def SWC(zh,rw,LJ,data=[]):
        global res
        wide=CeLiang.CeLiang(res,zh,rw).widen
        return [[zh,round(wide[0]+LJ,3)],[zh,0],[zh,round(wide[1]-LJ,3)]]
        

    
    # zh 桩号，rw 道路宽度，LJ 路肩宽度    
    def Other(zh,rw,LJ,data=[]):
        global res
        wide=CeLiang.CeLiang(res,zh,rw).widen
        return [[zh,wide[0]],[zh,0],[zh,wide[1]]]

    '''
    project:工程项目名称
    gongcheng:工程名称
    zhuanghao：桩号位置
    gongxu:工序
    data:数据[[SZ1,2.503,'','','','',2074.251,''],[ZD1,2.768,'',0.928,'','','','']]
    path:路径

    '''
    def write_gaocheng(project,gongcheng,zhuanghao,gongxu,data,path):
        num=len(data)
        page=(num-1)//17
        for p in range(0,page+1):
            wb=xl.load_workbook(muban+'\\高程.xlsx')
            ws=wb["sheet"]
            ws['A1']=project
            ws['A6']=f'工程名称：{gongcheng}'
            ws['D6']=f'{zhuanghao}{gongxu}'
            i=10
            w=p*17+17 if num>(p*17+17) else num
            for r in data[p*17:w]:
                j=1                  
                for c in r:
                   ws.cell(row=i, column=j,value =c)    
                   j=j+1
                i=i+1
            wb.save(path+f'\\{p}、{zhuanghao}{gongxu}-'+'高程检测.xlsx')
            wb.close()

    def write_record(project,gongcheng,zhuanghao,gongxu,header,data,path,No=1):
        num=len(data)
        pg=(num-1)//17
        yigao=round(random.uniform(1.2,1.6),3)
        for p in range(1,pg+2):            
            wb=xl.load_workbook(muban+'\\全站仪放线记录表.xlsx')
            ws=wb['sheet']
            ws['A1']=project
            ws['A5']=f'工程名称：{gongcheng}'
            ws['J5']=f'{zhuanghao}{gongxu}'
            ws['B7']=header[0][0] #测点编号
            ws['E7']=header[0][1] #测点X
            ws['E8']=header[0][2] #测点Y
            ws['E9']=header[0][3] #测点H
            ws['H7']=header[1][0] #后视点编号
            ws['K7']=header[1][1] #后视点X
            ws['K8']=header[1][2] #后视点Y
            ws['K9']=header[1][3] #后视点H
            ws['O7']=header[2]    #后视点距离
            ws['O9']=header[3]    #后视点方位角
            ws['T9']=yigao        #仪高
            i=12 
            w=(p-1)*16+16 if num>((p-1)*16+16) else num
            for r in data[(p-1)*17:w]:
                print('data:',r)
                j=1
                for c in r:
                    ws.cell(row=i, column=j,value =c)    
                    j=j+2
                i=i+1
            wb.save(path+f'\\{No}.{p}、{gongxu}-'+'全站仪放线记录表.xlsx')
            wb.close()
            

    def write_pingmianweizhi(project,gongcheng,zhuanghao,gongxu,header,data,path,No=1):
        num=len(data)

        pg=(num-1)//16
        yigao=round(random.uniform(1.2,1.6),3)
        
        for p in range(1,pg+2):        
            wb=xl.load_workbook(muban+'\\全站仪平面位置检测表.xlsx')
            ws=wb['sheet']
            ws['A1']=project
            ws['A5']=f'工程名称：{gongcheng}'
            ws['J5']=f'{zhuanghao}{gongxu}'
            ws['B7']=header[0][0] #测点编号
            ws['E7']=header[0][1] #测点X
            ws['E8']=header[0][2] #测点Y
            ws['E9']=header[0][3] #测点H
            ws['H7']=header[1][0] #后视点编号
            ws['K7']=header[1][1] #后视点X
            ws['K8']=header[1][2] #后视点Y
            ws['K9']=header[1][3] #后视点H
            ws['O7']=header[2]    #后视点距离
            ws['O9']=header[3]    #后视点方位角
            ws['T9']=yigao        #仪高
        #######################
            i=12         
            w=(p-1)*16+16 if num>((p-1)*16+16) else num
            for r in data[(p-1)*16:w]:
                j=1
             
                for c in r:
                    ws.cell(row=i, column=j,value =c)
                    j=j+2
                i=i+1
        ###################
            wb.save(path+f'\\{No}.{p}、{gongxu}-'+'平面位置检测.xlsx')
            wb.close()


############################################################################            
class Tool:

    def zhuanhuan(path,dest_path):
        pythoncom.CoInitialize()
        word= wc.Dispatch("Word.application")
        word.Visible =0
        word.DisplayAlerts =0
        excel=wc.gencache.EnsureDispatch('Excel.Application')
        for file in os.listdir(path):
           
            if os.path.splitext(file)[1] in ['.doc','.docx']:
                print("文件名",file)
                (file_path, temp_file_name) = os.path.split(file)
                (short_name, extension) = os.path.splitext(temp_file_name)
                doc = word.Documents.Open(path +'\\'+ file)
                doc.SaveAs(dest_path +'\\'+ short_name + ".docx", 16)
                doc.Close()
            elif os.path.splitext(file)[1] in ['.xls','.xlsx']:
                print("文件名",file)
                (file_path, temp_file_name) = os.path.split(file)
                (short_name, extension) = os.path.splitext(temp_file_name)
                wb = excel.Workbooks.Open(path +'\\'+ file)
                wb.SaveAs(dest_path +'\\'+ short_name + ".xlsx", 51)
                wb.Close()  
            
        word.Quit()
        excel.Quit()
        pythoncom.CoUninitialize()



    def replace_word(old_info, new_info,document):
        
        for paragraph in document.paragraphs:
            for run in paragraph.runs:
                if run.text and old_info in run.text:
                    rt=run.text.replace(old_info, new_info)
                    run.text=rt

        for table in document.tables:
            for row in table.rows:
                 for cell in row.cells:
                     for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            if run.text and old_info in run.text:
                                rt=run.text.replace(old_info, new_info)
                                run.text=rt

    def replace_excel(old,new,wb):    
        for ws in wb:
            for row in ws:
                for cell in row:
                    if cell.value is not None and old  in  str(cell.value) :
                        cell.value=cell.value.replace(old,new)


    '''
    打印文件
    '''
    def Printer(filename):
        open(filename, "r")
        win32api.ShellExecute(
            0,
            "print",
            filename,
            '/d:"%s"' % win32print.GetDefaultPrinter(),
            ".",
            0
        )
##############################################################
       
class APP:
    global Page

    global resources
    Page=None

    ENGS={'土方工程':'TF','涵洞工程':'HD','挡墙工程':'DQ_PJ','排水工程':'BG','路基工程':'JPSS','路面工程':'SWC','其他工程':'Other'}
    resources={'会东县人民村通村公路工程':'RMC','会东县营盘村通村公路工程':'YPC','会东县县道XW20线姜香路姜州至龙树段升级改造工程':'JXL','会东县县道XW21线新会路新街至铅锌镇段升级改造工程':'XHL','会东县县道XW22线会淌路升级改造工程':'HTL'}
    

    #########初始化###################################
    def __init__(self):
        global root
        global title
        global image
        global pic
        global bground
        global bg_name
        global Page
        global project
        global engineering
        global res
        

        res=None

        root=Tk()
        root.title('小工具 V1.0')
        root.geometry('1280x760')
        root.attributes("-alpha", 0.98)
        root.iconbitmap(RES+'\\'+'favicon.ico')


        mbar=Menu(root)
        menu_main=Menu(mbar,tearoff=0)
        menu_main.add_command(label='主页',command=self.MainPage)
        menu_main.add_command(label='批量替换',command=self.view_rep_SP)
        menu_main.add_command(label='文件夹内批量替换',command=self.view_Normal)
        menu_main.add_command(label='文档转换',command=self.view_exchange)
        menu_main.add_command(label='批量打印',command=self.Printer)

        mbar.add_cascade(label='主菜单',menu=menu_main)

        Pro=Menu(mbar,tearoff=0)
        project=StringVar()
        Pro.add_radiobutton(label='人民村路' ,variable=project,value='会东县人民村通村公路工程',command=self.SETTING)
        Pro.add_radiobutton(label='姜香路' ,variable=project,value='会东县县道XW20线姜香路姜州至龙树段升级改造工程',command=self.SETTING)
        Pro.add_radiobutton(label='新会路' ,variable=project,value='会东县县道XW21线新会路新街至铅锌镇段升级改造工程',command=self.SETTING)
        Pro.add_radiobutton(label='营盘村路',variable=project,value='会东县营盘村通村公路工程',command=self.SETTING)
        Pro.add_radiobutton(label='会淌路' ,variable=project,value='会东县县道XW22线会淌路升级改造工程',command=self.SETTING)
        Pro.add_radiobutton(label='其他' ,variable=project,value='other',command=self.SETTING)
        mbar.add_cascade(label='项目名称',menu=Pro)

        Class=Menu(mbar,tearoff=0)

        engineering=StringVar()

        Class.add_radiobutton(label='土方工程',variable=engineering ,value='土方工程',command=self.Deal)
        Class.add_radiobutton(label='涵洞工程',variable=engineering ,value='涵洞工程',command=self.Deal)
        Class.add_radiobutton(label='挡墙工程',variable=engineering ,value='挡墙工程',command=self.Deal)
        Class.add_radiobutton(label='排水工程',variable=engineering ,value='排水工程',command=self.Deal)
        Class.add_radiobutton(label='路肩工程',variable=engineering ,value='路肩工程',command=self.Deal)
        Class.add_radiobutton(label='路基工程',variable=engineering ,value='路基工程',command=self.Deal)
        Class.add_radiobutton(label='路面工程',variable=engineering ,value='路面工程',command=self.Deal)
        Class.add_radiobutton(label='其他工程',variable=engineering ,value='其他工程',command=self.Deal)
        mbar.add_cascade(label='类别',menu=Class)



        
        
        setting=Menu(mbar,tearoff=0)
        setting.add_command(label='换肤',command=self.skin)
        
        mbar.add_cascade(label='设置',menu=setting)
        mbar.add_command(label='关于')

        mbar.add_command(label='退出',command=root.quit)
        root.config(menu=mbar)
        bg_name=StringVar()
        bg_name.set(background+'\\fds.jpg')
        image=Image.open(f'{bg_name.get()}') if os.path.isfile(f'{bg_name.get()}') else None
        pic=ImageTk.PhotoImage(image) if os.path.isfile(f'{bg_name.get()}') else None
        msg=''
        bground=Label(root,text=msg,justify=LEFT,compound = CENTER,image=pic)
        bground.pack(side=LEFT)
       
        Page=PanedWindow(root,orient=VERTICAL)
        Page.place(x=50,y=60)

        self.Home()
        root.bind("<Configure>",self.background_image)
        

        root.mainloop()
    ###########背景调整################################
    def background_image(self,event):
        global root
        global image
        global pic
        global bground
        if os.path.isfile(f'{bg_name.get()}') :
            image=Image.open(f'{bg_name.get()}').resize((int(root.winfo_width()),int(root.winfo_height())))
            pic=ImageTk.PhotoImage(image)
            bground['image']=pic
    ############背景切换############################
    def image_change(self):
        global root
        global image
        global pic
        global bground
        image=Image.open(background+f'\\{bg_name.get()}').resize((int(root.winfo_width()),int(root.winfo_height())))
        pic=ImageTk.PhotoImage(image)
        bground['image']=pic
    ########################################
    def skin(self):
        global image
        global pic
        global bground
        bgs=[]
        for filename in os.listdir(background):
            bgs.append(filename)

        bg_name.set(f'{bgs[random.randint(0,len(bgs)-1)]}')
        image=Image.open(background+f'\\{bg_name.get()}').resize((int(root.winfo_width()),int(root.winfo_height())))
        pic=ImageTk.PhotoImage(image)
        bground['image']=pic
        ##################################################
    def SETTING(self):

        global project
        global res

        if project.get() in resources:
            res=RES+'\\'+resources[project.get()]
        else:
            res=None
        self.Home()

    def Deal(self):
        global engineering
        obj=SP
        self.fun=getattr(obj,self.ENGS[engineering.get()])

        self.MainPage()

    ########清除###########################################
    def Clear(self,Page:PanedWindow):
        
        if Page!=None:
            for  i in Page.panes():
                Page.forget(i)
                i=None
    ########主页#########################################
    def Home(self):
        global Page
        global project
        MyThread(self.Clear,Page)



        Page.add(Label(text='欢迎',font=('仿宋',20)))
        Page.add(Label(text='注：高程，平面位置需先设置，相关内容',font=('仿宋',13,'bold'),fg='red'))
        if project.get()=='会东县人民村通村公路工程':
            Page.add(Label(text='人民村路：0~6100路肩为0.25；6100~9645.5路肩为0.4，9645.5~10581.786为0.5；\n0~9645.5路面宽度为6.5，9645.5~10581.786为路面宽度7.5。',font=('仿宋',16),fg='brown'))
        
        
    def MainPage(self):
        global Page
        global project
        global engineering
        self.Clear(Page)
        if project.get()!='' and engineering.get()!='':
            Page.add(Button(text='放线记录',command=self.FX_view,height=5,width=40))
            Page.add(Button(text='平面位置检测',command=self.PMWZ_view,height=5,width=40))
            Page.add(Button(text='高程检测',command=self.HeightCheck,height=5,width=40))


    ########桩号替换#########################################
    def view_rep_SP(self):
        global root
        global Page        
        global values
        global HOME
        MyThread(self.Clear,Page)
        
        if Page not in globals():
            Page=PanedWindow(root,orient=VERTICAL)
            Page.place(x=60,y=60)
        Page.add(Label(text='特殊（桩号）批量替换工具',font=('仿宋 22')))
        pan1=PanedWindow()
        pan1.add(Label(text='需要处理的文件路径：'))
        path=StringVar()
        path.set('mode')
        pan1.add(Entry(textvariable=path))
        Page.add(pan1)
        pan2=PanedWindow()
        pan2.add(Label(text='需要处理的桩号文件：'))
        data=StringVar()
        data.set('rep.xlsx')
        pan2.add(Entry(textvariable=data))
        Page.add(pan2)
        values=[]
        PATH=None
        def deal():
            global PATH
            global SHEET
            try:
                if  os.path.isdir(path.get()):
                    PATH=path.get()  
                else:
                    raise Exception('未指定路径')
                if  os.path.isfile(data.get()):
                    dt=xl.load_workbook(data.get())
                    SHEET=list(dt.worksheets[0].rows)
                else:
                    raise Exception('未找到文件')
                row=list(dt.worksheets[0].rows)[0]
                if values==[]:
                    for cell in row:
                        pan=PanedWindow()
                        var1=StringVar()
                        pan.add(Label(text=f'替换{cell.column}（旧）：'))
                        var1.set(cell.value)
                        pan.add(Entry(textvariable=var1))
                        values.append(var1)
                        
                        Page.add(pan)
                    Page.add(Button(text='开始替换',command=self.replace_SP))
                dt.close()
            except Exception as err:

                messagebox.showerror('showerror', err)

        Page.add(Button(text='提交',command=deal))
            
    def replace_SP(self):
        global values
        global PATH
        global SHEET
        try:
            for row in SHEET:
                folder=' '.join([fo.value if fo.value!=None else 'X' for fo in row])
                if os.path.isdir(folder):
                    pass
                else:
                    os.makedirs(folder)

                for filename in os.listdir(PATH):
                    if os.path.splitext(filename)[1] in ['.docx']:
                        doc1=Document(f'{PATH}\\{filename}')
                        for cell in row:
                            VALUE=cell.value if cell.value!=None else ' '
                            Tool.replace_word(values[cell.column-1].get(),VALUE,doc1,)
                        doc1.save(f'{folder}\\{filename}')
                        
                    elif os.path.splitext(filename)[1] in ['.xlsx']:
                        wb=xl.load_workbook(f'{PATH}\\{filename}')
                        for cell in row:
                            VALUE=cell.value if cell.value!=None else ' '
                            Tool.replace_excel(values[cell.column-1].get(),VALUE,wb)
                        wb.save(f'{folder}\\{filename}')
                        wb.close()
                        
        except Exception as err:
            messagebox.showerror('错误', err)

        else:
            messagebox.showinfo('信息', '成功')            
    #########普通批量 替换###########################################
    def view_Normal(self):
        global Page
        global PATH
        global values
        MyThread(self.Clear,Page)
        Page.add(Label(text='指定路径批量替换',font=('仿宋 22')))
        opt1=PanedWindow()
        opt1.add(Label(text='需要处理的文件夹路径：'))

        PATH=StringVar()
        opt1.add(Entry(textvariable=PATH))
        
        Page.add(opt1)
        values=[]

        try:

            def add():
                opt2=PanedWindow()
                opt2.add(Label(text='旧'))
                old=StringVar()
                opt2.add(Entry(textvariable=old))
                opt2.add(Label(text='新'))
                new=StringVar()
                opt2.add(Entry(textvariable=new))
                values.append([old,new])
                INDEX=values.index([old,new])
                def close(pan,index):
                    Page.remove(pan)
                    values.pop(index)
                opt2.add(Button(text='×',command=lambda:close(opt2,INDEX)))
                Page.add(opt2)
        except Exception as e:
            messagebox.showerror('错误', e)


        opt3=PanedWindow()
        
        opt3.add(Button(text='增加',command=add))
        Page.add(opt3)

        Page.add(Button(text='开始替换',command=self.replace_N))
        pass
    def replace_N(self):
        global values
        global PATH

        try:
            for filename in os.listdir(PATH.get()):

                if os.path.splitext(filename)[1] in ['.docx']:
                    doc1=Document(f'{PATH.get()}\\{filename}')
                    for row in values:
                        Tool.replace_word(row[0].get(),row[1].get(),doc1)

                    doc1.save(f'{PATH.get()}\\{filename}')
                elif os.path.splitext(filename)[1] in ['.xlsx']:
                    wb=xl.load_workbook(f'{PATH.get()}\\{filename}')
                    for row in values:
                        Tool.replace_excel(row[0].get(),row[1].get(),wb)
                    wb.save(f'{PATH.get()}\\{filename}')
        except Exception as e:
            messagebox.showerror('错误', e)
        else:
            messagebox.showinfo('信息', '成功')
    ##########格式转换#############################################
    def view_exchange(self):
        global Page
        global PATH
        global PATH2        
        self.Clear(Page)
        Page.add(Label(text='Word、Excel 转换成 .docx,.xlsx',font=('仿宋 22')))
        opt1=PanedWindow()
        opt1.add(Label(text='需要处理的文件夹路径：'))
        PATH=StringVar()
        opt1.add(Entry(textvariable=PATH))        
        Page.add(opt1)
        opt2=PanedWindow()
        opt2.add(Label(text='存储路径：'))
        PATH2=StringVar()
        opt2.add(Entry(textvariable=PATH2)) 
        Page.add(opt2)
        Page.add(Button(text='开始转换',command=self.exchange))
    def exchange(self):
        global PATH
        global PATH2
        if os.path.isdir(PATH.get()):
            if os.path.isdir(PATH2.get()):
                path=PATH2.get()
            else:
                os.makedirs(PATH2.get())
                path=PATH2.get()

            Tool.zhuanhuan(PATH.get(),path)
    #############批量打印#############################################
    def Printer(self):
        global Page
        global opt_md
        global show_md
        global show_data
        MyThread(self.Clear,Page)
        show_md=False
        Page.add(Label(text='批量打印',font=('仿宋',18,'bold'),fg='purple'))
        opt1=PanedWindow()
        opt1.add(Label(text='目标文件夹'))
        PATH=StringVar()
        opt1.add(Entry(textvariable=PATH))
        opt_list=None
        opt1.add(Button(text='获取文件',command=lambda:self.getFiles(opt_list,PATH)))
        opt1.add(Button(text='获取文件夹',command=lambda:self.getFolders(opt_list,PATH)))
        Page.add(opt1)
        opt_list=PanedWindow(height=300,width=500)
        Page.add(opt_list)
        show_data=False

        opt2=PanedWindow()
        opt2.add(Label(text='打印份数：'))
        num=IntVar()
        num.set(1)
        opt2.add(Entry(textvariable=num))
        opt2.add(Label(text='每份结束等待（秒）：'))
        sl=IntVar()
        sl.set(3)
        opt2.add(Entry(textvariable=sl))
        opt2.add(Button(text='打印',command=lambda:self.PRINT(num,sl)))
        Page.add(opt2)
        opt_md=PanedWindow(height=30)
        Page.add(opt_md)
    def getFiles(self,top,path):
        global tree
        global show_data
        global mode
        if show_data==False:
            mode='file'
            tree=ttk.Treeview(top,show="headings")
            s=ttk.Style()
            s.theme_use('default')
            tree['columns']=['路径','文件名']
            tree.column("路径",width=100,anchor="center")
            tree.column("文件名",width=100,anchor="center")
            tree.heading("路径",text="路径")
            tree.heading("文件名",text="文件名")
            i=1
            files=os.listdir(path.get())        
            for filename in natsorted(files):
                if os.path.isfile(path.get()+'\\'+filename):
                    tree.insert('',i,values=(path.get(),filename))
                    i=i+1

            tree.bind("<Delete>",self.Del)
                 
            top.add(tree)
            show_data=True
    def getFolders(self,top,path):
        global tree
        global show_data
        global mode
        if show_data==False:
            mode='folder'
            tree=ttk.Treeview(top,show="headings")
            s=ttk.Style()
            s.theme_use('default')
            tree['columns']=['路径','子文件夹名']
            tree.column("路径",width=100,anchor="center")
            tree.column("子文件夹名",width=100,anchor="center")
            tree.heading("路径",text="路径")
            tree.heading("子文件夹名",text="子文件夹名")
            i=0
            files=os.listdir(path.get())

            for filename in natsorted(files):
                if os.path.isdir(path.get()+'\\'+filename):
                    k=len(natsorted(os.listdir(f'{path.get()}\\{filename}')))
                    tree.insert('','end',f'ID{i}',values=(path.get(),filename))
                    j=1
                    for file in natsorted(os.listdir(f'{path.get()}\\{filename}')):
                        tree.insert(f'ID{i}','end',f'print{i}.{j}',values=(f'{path.get()}\\{filename}',file))
                        j=j+1
                    i=i+1

            tree.bind("<Delete>",self.Del)
            
            top.add(tree)
            show_data=True
        else:
            i=0
            files=os.listdir(path.get())

            for filename in natsorted(files):
                if os.path.isdir(path.get()+'\\'+filename):
                    k=len(natsorted(os.listdir(f'{path.get()}\\{filename}')))
                    tree.insert('','end',f'ID{i}',values=(path.get(),filename))
                    j=1
                    for file in natsorted(os.listdir(f'{path.get()}\\{filename}')):
                        tree.insert(f'ID{i}','end',f'print{i}.{j}',values=(f'{path.get()}\\{filename}',file))
                        j=j+1
                    i=i+1            
    def PRINT(self,num,sl):
        global tree
        global mode

        try:

            SL=int(sl.get()) if int(sl.get())>3 else 3
            NUM=int(num.get())

            if mode=='folder':
                for item in tree.get_children():
                    for n in range(NUM):
                        print(f'开始打印第{n+1}份')
                        for file in tree.get_children(item):
                            filename='\\'.join(tree.item(file,'values'))
                            Tool.Printer(filename)
                            time.sleep(SL/2)
                        time.sleep(SL)
            elif mode=='file':
                for n in range(NUM):
                    print(f'开始打印第{n+1}份')
                    for file in tree.get_children():
                        filename='\\'.join(tree.item(file,'values'))
                        
                        Tool.Printer(filename)
                        time.sleep(SL)
        except Exception as e:
            messagebox.showerror('错误',f'打印失败：{e}')
        else:
            messagebox.showinfo('信息','打印结束')

    ###########平面位置###################################################
    def View(self,title,JC,tb_hd,font,rowHigh,cmd):
        global project
        global Page
        global pulldown
        global engineering
        global Pane3
        global side
        global show_md
        global Pane_mod
        global tree
        global JianCe
        JianCe=JC

        show_md=False

        Pane1=PanedWindow(height=rowHigh*2)
        Pane1.add(Label(text=title,font=font))
        Page.add(Pane1)
        Pane2=PanedWindow(height=rowHigh)
        Pane2.add(Label(text='工程项目名称'))
        Pane2.add(Entry(textvariable=project,width=40))
        Pane2.add(Label(text='工程名称'))
        LIST=list(self.ENGS.keys())
        LIST.insert(0,'选择')
        engineering=StringVar()
        pulldown=ttk.Combobox(values=LIST,textvariable=engineering)
        pulldown.current(0)
        pulldown['state']='readonly'
        Pane3=None
        pulldown.bind("<<ComboboxSelected>>",self.showSide)
        Pane2.add(pulldown)
        Page.add(Pane2)
        Pane3=PanedWindow(height=rowHigh)
        Pane3.add(Label(text='工程部位'))

        ZH1=DoubleVar()
        ZH2=DoubleVar()
        Ping=IntVar()
        RW=DoubleVar()
        LJ=DoubleVar()
        side=StringVar()
        high=DoubleVar()
        RW.set(6.5)
        LJ.set(0.25)

        Pane3.add(Entry(textvariable=ZH1))
        Pane3.add(Label(text='--'))
        Pane3.add(Entry(textvariable=ZH2))

        Pane3.add(Label(text='相对高度：'))
        Pane3.add(Entry(textvariable=high))

        
        Page.add(Pane3)

 
        Pane4=PanedWindow(height=rowHigh)
        Pane4.add(Label(text='工序'))
        gongxu=ttk.Combobox(values=['基坑开挖前','基坑开挖后','基坑底','基层顶','墙身顶','回填第i层'])
        Pane4.add(gongxu)
        Pane4.add(Label(text='检测频率'))
        Pane4.add(Entry(textvariable=Ping))
        Pane4.add(Label(text='路面宽度'))
        Pane4.add(Entry(textvariable=RW))
        Pane4.add(Label(text='路肩宽度'))
        Pane4.add(Entry(textvariable=LJ))
        Page.add(Pane4)
        Pane5=PanedWindow(height=rowHigh)
        Pane_data=None

        Pane5.add(Button(text='自动获取',width=30,command=lambda:MyThread(self.Tree,Pane_data,tb_hd,ZH1,ZH2,side,Ping,RW,LJ,high,pulldown,gongxu)))
        Pane5.add(Button(text='修改计算',width=30,command=lambda:MyThread(self.calculate,tree,high,RW)))
        Page.add(Pane5)
        Page.add(Label(text='++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++'))
        Pane_mod=PanedWindow(height=rowHigh)
        Page.add(Pane_mod)
        Pane_data=PanedWindow(height=350)
        Page.add(Pane_data)
        Pane_save=PanedWindow()
        Pane_save.add(Label(text='保存位置：'))
        PATH=StringVar()
        Pane_save.add(Entry(textvariable=PATH,width=30))
        Pane_save.add(Button(text='保存(SAVE)',command=lambda:MyThread(cmd,tree,project,pulldown,ZH1,ZH2,side,gongxu,PATH)))

        Page.add(Pane_save)
        pass

    def showSide(self,event):
        global side
        global pulldown
        global Pane3

        if pulldown.get() in list(self.ENGS.keys())[3:11]:
            if side.get()=='':
                side=StringVar()
                side.set('左侧')
                Pane3.add(Radiobutton(text='左侧' ,variable=side,value='左侧'))
                Pane3.add(Radiobutton(text='右侧' ,variable=side,value='右侧'))
            pulldown['state']='readonly' 
        else:
            pulldown['state']='' if pulldown.get()=='其他' else 'readonly'
            side.set('')
            for pan in Pane3.panes():
                if 'radiobutton' in pan.string:
                    Pane3.forget(pan)


    def FX_view(self):
        global Page
        global JianCe
        MyThread(self.Clear,Page)
        JianCe='pmwz'

        self.View('放线记录','放线记录',['桩号','偏距','X','Y','计算方位角','计算距离'],('仿宋',20),25,self.gaochengDeal)

        
    ###########高程####################################################
    def HeightCheck(self):
        global Page
        global JianCe
        MyThread(self.Clear,Page)
        JianCe='pmwz'

        self.View('高程检测','高程检测',['桩号','偏距','X','Y'],('仿宋',20),25,self.gaochengDeal)


    def gaochengDeal(self,tree,project,side,engineering,zh1,zh2,gongxu,PATH):

        try:
            ZH1=round(float(zh1.get()),3)
            ZH2=round(float(zh2.get()),3)
            path=f'{PATH.get()}\\{engineering.get()}\\{SP.mileageToStr(ZH1)}-{SP.mileageToStr(ZH2)}{side.get()}{gongxu.get()}'
            if os.path.isdir(path):
                pass
            else:
                os.makedirs(path)

            data=[]

            BM=None
            Hsx=None
            dest=None
            zhq=None
            i=0

            ZHS=tree.get_children()
            for item in ZHS:
                row=tree.item(item,'values')

                zh=round(float(row[0]),3)
                pianju=row[1]
                HS=round(float(row[2]),3)
                pc=random.randint(-20,20)
                if BM!=OPTION.getBM(zh):
                
                    BM=OPTION.getBM(zh)
                    
                    
                    H=BM[2]-HS
                    L=abs(zh-BM[0])
                    Hsx=None
                    i=0
                    back=SP.ZD(BM,H,abs(L),data,i,Hsx)
                    Hsx=back[0]
                    i=back[2]
                    dest=back[1]
                    zhq=zh
                    data.append([f'{SP.mileageToStr(zh)},{pianju}','','',round(Hsx-round(HS+pc/1000,3),3),'',round(HS+pc/1000,3),HS,pc])
                else:
                    HC=Hsx-HS
                    
                    L=zh-zhq
                    
                    if HC<0.6 or HC>4.8:
                        H=Hsx-HS-dest if HC> 4.8 else HC-2
                        back=SP.ZD(BM,H,abs(L),data,i,Hsx) if abs(L)!=0 else SP.ZD(BM,H,abs(L+2),data,i,Hsx)
                        print(back)
                        Hsx=back[0]
                        i=back[2]
                        zhq=zh        
                    data.append([f'{SP.mileageToStr(zh)},{pianju}','','',round(Hsx-round(HS+pc/1000,3),3),'',round(HS+pc/1000,3),HS,pc])
            print(path)
            SP.write_gaocheng(project.get(),engineering.get(),f'{SP.mileageToStr(round(float(zh1.get()),3))}-{SP.mileageToStr(round(float(zh2.get()),3))}',f'({side.get()}){gongxu.get()}',data,path)
        
            
        except Exception as err:
            messagebox.showerror('错误信息',f'发生错误,{err}，执行失败！！')
        else:

            messagebox.showinfo('信息', '成功')
            self.HeightCheck()

    def PMWZ_view(self):
        global Page
        global JianCe
        self.Clear(Page)
        

        self.View('平面位置检测','平面位置检测',['桩号','偏距','X','Y'],('仿宋',20),25,self.pmwzDeal)

    def pmwzDeal(self,tree,project,engineering,zh1,zh2,side,gongxu,PATH):

        try:
            ZH1=round(float(zh1.get()),3)
            ZH2=round(float(zh2.get()),3)

            path=f'{PATH.get()}\\{engineering.get()}\\{SP.mileageToStr(ZH1)}-{SP.mileageToStr(ZH2)}({side.get()}){gongxu.get()}'
            if os.path.isdir(path):
                pass
            else:
                os.makedirs(path)
            print(path)
            HEAD=None
            data=[]
            No=1
            
            ZHS=tree.get_children()
            for item in ZHS:
                row =tree.item(item,'values')

                zh=round(float(row[0]),3)
                pianju=row[1]
                X=round(float(row[2]),4)
                Y=round(float(row[3]),4)
                pc=int(round(((50**2)/2)**0.5))
                px=random.randint(-pc,pc)
                py=random.randint(-pc,pc)
                ps=round((px**2+py**2)**0.5)
                if HEAD!=SP.HEADER(OPTION.getCtr(zh)):

                    if data!=[]:
                        print('data',data)

                        SP.write_pingmianweizhi(project.get(),engineering.get(),f'{SP.mileageToStr(ZH1)}-{SP.mileageToStr(ZH2)}',f'({side.get()}){gongxu.get()}',HEAD,data,path,No)
                        No=No+1
                        data=[]
                    
                    data.append([f'{SP.mileageToStr(zh)} {pianju}',X,Y,round(X+px/1000,4),round(Y+py/1000,4),px,py,ps])
                else:
                    
                    data.append([f'{SP.mileageToStr(zh)} {pianju}',X,Y,round(X+px/1000,4),round(Y+py/1000,4),px,py,ps])
                HEAD=SP.HEADER(OPTION.getCtr(zh))

            if data!=[]:
                print('data',data)
                SP.write_pingmianweizhi(project.get(),engineering.get(),f'{SP.mileageToStr(ZH1)}-{SP.mileageToStr(ZH2)}',f'({side.get()}){gongxu.get()}',HEAD,data,path,No)


        except Exception as e:
            messagebox.showerror('错误信息',f'发生错误,{e}，执行失败！！')
        else:
            messagebox.showinfo('信息', '成功')
            


    ####################################################################

    def Tree(self,top,header,zh1,zh2,side,ping,rw,lj,high,eng,gongxu):
        global tree
        global JianCe
        global res

        print('桩号1',zh1.get())
        print('桩号2',zh2.get())
        print('频率',ping.get())
        print('高差',high.get())
        print('工程',eng.get())
        print('gongxu',gongxu.get())
        print('方法',self.fun)

        if 'tree' in globals():

            num=1 if ping.get()==0 else int((zh2.get()-zh1.get())//ping.get())
            CDS=[round(zh1.get()+ping.get()*i+ping.get()*random.uniform(0.3,0.8),3) for i in range(num)]
            
            #zh,rw,LJ,data=[]
            data=[high.get(),side.get(),eng.get(),gongxu.get()]
            NO=0
            index=0
            for zh in CDS:
                dd=self.fun(zh,rw.get(),lj.get(),data)
                
                if dd!=None:
                    k=len(dd)
                    i=0
                    for d in dd:
                        height=CeLiang.CeLiang(res,d[0],rw.get()).Height(d[1])
                        point=CeLiang.CeLiang(res,d[0],rw.get()).Point(d[1])
                        if d[1]==0:
                            pj='中'
                        else:
                            pj=f'左,{-d[1]}' if d[1]<0 else f'右,{d[1]}'
                        values=[d[0],pj,f'{point[0]:.4f}',f'{point[1]:.4f}']
                        NO=index*k+i
                        tree.item(f'{NO}',values=values) if tree.exists(f'{index*k+i}') else tree.insert('','end',f'{index*k+i}',values=values)
                        i=i+1
                    index=index+1
            for item in tree.get_children()[NO+1:]:
                if tree.exists(item):
                    tree.delete(item)


        else:
            tree=ttk.Treeview(show='headings')#show='headings'

            style=ttk.Style()
            style.theme_use('default')
            
            tree["columns"] = header
            for head in header:
                tree.column(f"{head}", width=80,anchor="center")
                tree.heading(f"{head}", text=f"{head}")

            num=1 if ping.get()==0 else int((zh2.get()-zh1.get())//ping.get())
            CDS=[round(zh1.get()+ping.get()*i+ping.get()*random.uniform(0.3,0.8),3) for i in range(num)]
            
            #zh,rw,LJ,data=[]
            data=[high.get(),side.get(),eng.get(),gongxu.get()]
            
            index=0
            for zh in CDS:
                dd=self.fun(zh,rw.get(),lj.get(),data)
                
                if dd!=None:
                    k=len(dd)
                    i=0
                    for d in dd:
                        height=CeLiang.CeLiang(res,d[0],rw.get()).Height(d[1])
                        point=CeLiang.CeLiang(res,d[0],rw.get()).Point(d[1])
                        if d[1]==0:
                            pj='中'
                        else:
                            pj=f'左,{-d[1]}' if d[1]<0 else f'右,{d[1]}'
                        values=[d[0],pj,f'{point[0]:.4f}',f'{point[1]:.4f}']
                        
                        tree.insert('','end',f'{index*k+i}',values=values)
                        
                        i=i+1
                    index=index+1
            
            
            tree.bind("<Delete>",self.Del)
            tree.bind("<Double-1>",self.edit) 

            top.add(tree)




    ############################################################
    def Del(self,event):
        global tree
        
        for item in tree.selection():
            if tree.exists(item):
                tree.delete(item)


    def edit(self,event):

        global tree
        global mods
        global Pane_mod
        global show_md
        for item in tree.selection():
            #item = I001
            print(item)
            item_text = tree.item(item, "values")

            def save(Item):
                global show_md
                values=[s.get() for s in mods]
                tree.item(Item, text="", values=values)
                tree.update()
                self.Clear(Pane_mod)
                show_md=False
                print(Item)
                

            if show_md==False:
                mods=[]
                lbs=tree['columns']
                i=0
                for var in item_text:
                    tem=StringVar()
                    tem.set(var)
                    entryedit = Entry(textvariable=tem,width=10)
                    Pane_mod.add(Label(text=f'{lbs[i]}：'))
                    Pane_mod.add(entryedit)
                    mods.append(tem)
                    i=i+1
                Pane_mod.add(Button(text='保存',command=lambda:save(item)))
                show_md=True
            

    def calculate(self,tree,RW,ReH):
        global JianCe
        
        global res
        try:

            if RW.get()!='' and ReH.get()!='':
                rw=round(float(RW.get()),3)
                h0=round(float(ReH.get()),3)
                for item in tree.get_children():
                    values=tree.item(item,'values')
                    zh=round(float(values[0]),3)
                    if '右' in values[1].split(','):
                        pj=round(float(values[1].split(',')[1]),3)
                    elif '左' in values[1].split(','):
                        pj=round(-float(values[1].split(',')[1]),3)
                    else:
                        pj=0
                    H=round(CeLiang.CeLiang(res,zh,rw).Height(pj)[1]+h0,3)
                    point=CeLiang.CeLiang(res,zh,rw).Point(pj)

                    data=['' for i in range(len(tree['columns']))]
                    data[0]=values[0]
                    data[1]=values[1]
                    print(JianCe)
                    if JianCe=='高程检测':
                        data[2]=H
                    elif JianCe=='平面位置检测':
                        data[2]=f'{point[0]:.4f}'
                        data[3]=f'{point[1]:.4f}'

                    tree.item(item, text="", values=data)
                    tree.update()
            else:
                raise Exception('请检查相关参数是否输入完整')
        except Exception as e:
            messagebox.showerror('错误', f'计算失败,发生错误：{e}')

    ################################################################
    @staticmethod
    def thread_it(func, *args):
        t = threading.Thread(target=func, args=args) 
        t.setDaemon(True)   # 守护--就算主界面关闭，线程也会留守后台运行（不对!）
        t.start()           # 启动
        # t.join()          # 阻塞--会卡死界面！


APP()


    
